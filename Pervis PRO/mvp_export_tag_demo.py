#!/usr/bin/env python3
"""
PreVis PRO 增强导出和标签管理系统 - MVP演示脚本
演示所有核心功能：文档导出、图片导出、标签管理、向量分析
"""

import sys
import os
import json
from pathlib import Path
from datetime import datetime

# 添加backend到路径
sys.path.insert(0, str(Path(__file__).parent / "backend"))

print("=" * 80)
print("PreVis PRO 增强导出和标签管理系统 - MVP演示")
print("=" * 80)
print()

# ============================================================================
# 第一部分：数据库初始化和数据准备
# ============================================================================

print("📊 第一部分：数据库初始化")
print("-" * 80)

try:
    from database import init_database, SessionLocal, Project, Beat, Asset, TagHierarchy, AssetTag
    
    # 初始化数据库
    print("✓ 正在初始化数据库...")
    init_database()
    print("✓ 数据库初始化完成")
    
    # 创建数据库会话
    db = SessionLocal()
    print("✓ 数据库连接成功")
    
except Exception as e:
    print(f"✗ 数据库初始化失败: {e}")
    sys.exit(1)

print()

# ============================================================================
# 第二部分：加载Cyberpunk演示项目数据
# ============================================================================

print("🎬 第二部分：加载Cyberpunk演示项目")
print("-" * 80)

demo_project_path = Path("demo_projects/cyberpunk_trailer")

try:
    # 加载项目数据
    with open(demo_project_path / "project.json", "r", encoding="utf-8") as f:
        project_data = json.load(f)
    print(f"✓ 加载项目: {project_data['title']}")
    
    # 加载Beat数据
    with open(demo_project_path / "beats.json", "r", encoding="utf-8") as f:
        beats_data = json.load(f)
    print(f"✓ 加载 {len(beats_data['beats'])} 个Beat")
    
    # 加载标签数据
    with open(demo_project_path / "tags.json", "r", encoding="utf-8") as f:
        tags_data = json.load(f)
    print(f"✓ 加载标签系统 (版本 {tags_data['tag_system_version']})")
    
except Exception as e:
    print(f"✗ 加载演示数据失败: {e}")
    sys.exit(1)

print()
print(f"项目信息:")
print(f"  - 标题: {project_data['title']}")
print(f"  - 类型: {project_data['type']}")
print(f"  - 类别: {project_data['genre']}")
print(f"  - Beat数量: {len(beats_data['beats'])}")
print(f"  - 标签类别: {len(tags_data['tag_categories'])}")
print()

# ============================================================================
# 第三部分：文档导出功能演示
# ============================================================================

print("📄 第三部分：文档导出功能")
print("-" * 80)

print("\n[1] DOCX格式导出")
print("-" * 40)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 创建DOCX文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(project_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加元数据
    doc.add_paragraph(f"类型: {project_data['genre']}")
    doc.add_paragraph(f"创建时间: {project_data['created_at']}")
    doc.add_paragraph(f"Beat数量: {len(beats_data['beats'])}")
    doc.add_paragraph()
    
    # 添加每个Beat
    for beat in beats_data['beats']:
        # Beat标题
        beat_heading = doc.add_heading(f"Beat {beat['sequence']}: {beat['title']}", level=1)
        
        # Beat内容
        doc.add_paragraph(f"场景类型: {beat['scene_type']}")
        doc.add_paragraph(f"预估时长: {beat['duration_estimate']}")
        doc.add_paragraph()
        
        # 内容描述
        doc.add_heading("场景描述", level=2)
        content = beat['content']
        doc.add_paragraph(f"地点: {content['location']}")
        doc.add_paragraph(f"时间: {content['time']}")
        doc.add_paragraph(f"动作: {content['action']}")
        doc.add_paragraph(f"情绪: {content['mood']}")
        doc.add_paragraph()
        
        # 视觉元素
        doc.add_heading("视觉元素", level=2)
        visual = beat['visual_elements']
        doc.add_paragraph(f"摄影机运动: {visual['camera_movement']}")
        doc.add_paragraph(f"光线: {visual['lighting']}")
        doc.add_paragraph(f"构图: {visual['composition']}")
        doc.add_paragraph(f"色彩: {visual['color_palette']}")
        doc.add_paragraph()
        
        # 标签
        doc.add_heading("标签", level=2)
        tags_text = ", ".join(beat['tags'])
        doc.add_paragraph(tags_text)
        doc.add_paragraph()
        
        # 分页
        if beat['sequence'] < len(beats_data['beats']):
            doc.add_page_break()
    
    # 保存文档
    output_dir = Path("exports")
    output_dir.mkdir(exist_ok=True)
    docx_path = output_dir / f"{project_data['id']}_script.docx"
    doc.save(str(docx_path))
    
    file_size = docx_path.stat().st_size
    print(f"✓ DOCX导出成功")
    print(f"  文件路径: {docx_path}")
    print(f"  文件大小: {file_size:,} 字节 ({file_size/1024:.1f} KB)")
    print(f"  页数: {len(beats_data['beats'])} 页")
    
except ImportError:
    print("✗ python-docx未安装，跳过DOCX导出")
    print("  安装命令: pip install python-docx")
except Exception as e:
    print(f"✗ DOCX导出失败: {e}")

print("\n[2] PDF格式导出")
print("-" * 40)

try:
    from weasyprint import HTML, CSS
    from jinja2 import Template
    
    # HTML模板
    html_template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: 'Microsoft YaHei', Arial, sans-serif;
                line-height: 1.6;
                color: #333;
            }
            h1 {
                color: #fbbf24;
                text-align: center;
                font-size: 32px;
                margin-bottom: 20px;
            }
            h2 {
                color: #3b82f6;
                font-size: 24px;
                margin-top: 30px;
                border-bottom: 2px solid #3b82f6;
                padding-bottom: 5px;
            }
            h3 {
                color: #10b981;
                font-size: 18px;
                margin-top: 20px;
            }
            .metadata {
                background: #f3f4f6;
                padding: 15px;
                border-radius: 5px;
                margin-bottom: 30px;
            }
            .beat {
                page-break-after: always;
                margin-bottom: 40px;
            }
            .beat:last-child {
                page-break-after: auto;
            }
            .tags {
                background: #fef3c7;
                padding: 10px;
                border-radius: 5px;
                margin-top: 10px;
            }
            .visual-elements {
                background: #dbeafe;
                padding: 10px;
                border-radius: 5px;
                margin-top: 10px;
            }
        </style>
    </head>
    <body>
        <h1>{{ title }}</h1>
        <div class="metadata">
            <p><strong>类型:</strong> {{ genre }}</p>
            <p><strong>创建时间:</strong> {{ created_at }}</p>
            <p><strong>Beat数量:</strong> {{ beat_count }}</p>
        </div>
        
        {% for beat in beats %}
        <div class="beat">
            <h2>Beat {{ beat.sequence }}: {{ beat.title }}</h2>
            <p><strong>场景类型:</strong> {{ beat.scene_type }}</p>
            <p><strong>预估时长:</strong> {{ beat.duration_estimate }}</p>
            
            <h3>场景描述</h3>
            <p><strong>地点:</strong> {{ beat.content.location }}</p>
            <p><strong>时间:</strong> {{ beat.content.time }}</p>
            <p><strong>动作:</strong> {{ beat.content.action }}</p>
            <p><strong>情绪:</strong> {{ beat.content.mood }}</p>
            
            <h3>视觉元素</h3>
            <div class="visual-elements">
                <p><strong>摄影机运动:</strong> {{ beat.visual_elements.camera_movement }}</p>
                <p><strong>光线:</strong> {{ beat.visual_elements.lighting }}</p>
                <p><strong>构图:</strong> {{ beat.visual_elements.composition }}</p>
                <p><strong>色彩:</strong> {{ beat.visual_elements.color_palette }}</p>
            </div>
            
            <h3>标签</h3>
            <div class="tags">
                {{ beat.tags | join(', ') }}
            </div>
        </div>
        {% endfor %}
    </body>
    </html>
    """
    
    # 渲染HTML
    template = Template(html_template)
    html_content = template.render(
        title=project_data['title'],
        genre=project_data['genre'],
        created_at=project_data['created_at'],
        beat_count=len(beats_data['beats']),
        beats=beats_data['beats']
    )
    
    # 生成PDF
    pdf_path = output_dir / f"{project_data['id']}_script.pdf"
    HTML(string=html_content).write_pdf(str(pdf_path))
    
    file_size = pdf_path.stat().st_size
    print(f"✓ PDF导出成功")
    print(f"  文件路径: {pdf_path}")
    print(f"  文件大小: {file_size:,} 字节 ({file_size/1024:.1f} KB)")
    
except ImportError:
    print("✗ WeasyPrint未安装，跳过PDF导出")
    print("  安装命令: pip install weasyprint")
except Exception as e:
    print(f"✗ PDF导出失败: {e}")

print()

# ============================================================================
# 第四部分：BeatBoard图片导出
# ============================================================================

print("🖼️  第四部分：BeatBoard图片导出")
print("-" * 80)

try:
    from PIL import Image, ImageDraw, ImageFont
    
    # 创建BeatBoard可视化
    width, height = 1920, 1080
    img = Image.new('RGB', (width, height), color='#1a1a1a')
    draw = ImageDraw.Draw(img)
    
    # 尝试加载字体
    try:
        title_font = ImageFont.truetype("msyh.ttc", 48)  # 微软雅黑
        heading_font = ImageFont.truetype("msyh.ttc", 32)
        body_font = ImageFont.truetype("msyh.ttc", 24)
        small_font = ImageFont.truetype("msyh.ttc", 18)
    except:
        title_font = ImageFont.load_default()
        heading_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
    
    # 绘制标题
    title_text = project_data['title']
    draw.text((width//2, 60), title_text, fill='#fbbf24', font=title_font, anchor='mm')
    
    # 绘制Beat卡片
    beat_width = 500
    beat_height = 280
    margin = 40
    cols = 3
    
    for i, beat in enumerate(beats_data['beats']):
        row = i // cols
        col = i % cols
        
        x = margin + col * (beat_width + margin)
        y = 150 + row * (beat_height + margin)
        
        # 绘制卡片背景
        draw.rectangle([x, y, x + beat_width, y + beat_height], fill='#2d2d2d', outline='#fbbf24', width=2)
        
        # Beat标题
        beat_title = f"Beat {beat['sequence']}: {beat['title']}"
        draw.text((x + 20, y + 20), beat_title, fill='#f3f4f6', font=heading_font)
        
        # 场景类型
        draw.text((x + 20, y + 70), f"类型: {beat['scene_type']}", fill='#d1d5db', font=body_font)
        
        # 时长
        draw.text((x + 20, y + 105), f"时长: {beat['duration_estimate']}", fill='#d1d5db', font=body_font)
        
        # 情绪强度
        intensity = beat['emotional_arc']['intensity']
        draw.text((x + 20, y + 140), f"情绪强度: {intensity}/10", fill='#fbbf24', font=body_font)
        
        # 标签（前3个）
        tags_text = ", ".join(beat['tags'][:3])
        if len(beat['tags']) > 3:
            tags_text += "..."
        draw.text((x + 20, y + 180), tags_text, fill='#9ca3af', font=small_font)
        
        # 情绪曲线指示器
        emotion_colors = {
            '平静': '#10b981',
            '紧张': '#f59e0b',
            '恐惧': '#ef4444',
            '压迫感': '#8b5cf6',
            '决心': '#3b82f6'
        }
        start_emotion = beat['emotional_arc']['start_emotion']
        color = emotion_colors.get(start_emotion, '#6b7280')
        draw.ellipse([x + beat_width - 60, y + 20, x + beat_width - 20, y + 60], fill=color)
    
    # 保存图片
    image_path = output_dir / f"{project_data['id']}_beatboard.png"
    img.save(str(image_path), 'PNG', quality=95)
    
    file_size = image_path.stat().st_size
    print(f"✓ BeatBoard图片导出成功")
    print(f"  文件路径: {image_path}")
    print(f"  文件大小: {file_size:,} 字节 ({file_size/1024:.1f} KB)")
    print(f"  分辨率: {width}x{height}")
    print(f"  格式: PNG")
    
except ImportError:
    print("✗ Pillow未安装，跳过图片导出")
    print("  安装命令: pip install Pillow")
except Exception as e:
    print(f"✗ 图片导出失败: {e}")

print()

# ============================================================================
# 第五部分：标签层级管理
# ============================================================================

print("🏷️  第五部分：标签层级管理")
print("-" * 80)

# 构建标签层级树
tag_hierarchy = {}
for category, scenes in tags_data['tag_categories'].items():
    category_name = category.replace('_tags', '')
    tag_hierarchy[category_name] = {}
    
    for scene, tags in scenes.items():
        tag_hierarchy[category_name][scene] = tags

print("\n标签层级结构:")
print("-" * 40)

for category, scenes in tag_hierarchy.items():
    print(f"\n📁 {category}")
    for scene, tags in scenes.items():
        print(f"  └─ {scene}")
        for tag in tags[:3]:  # 只显示前3个
            print(f"      • {tag}")
        if len(tags) > 3:
            print(f"      ... 还有 {len(tags) - 3} 个标签")

# 标签权重示例
print("\n\n标签权重调整示例:")
print("-" * 40)

sample_weights = {
    "城市": 0.95,
    "夜晚": 0.92,
    "追逐": 0.88,
    "紧张": 0.90,
    "霓虹": 0.85
}

print("\nBeat 1 (城市追逐) 的标签权重:")
for tag, weight in sample_weights.items():
    bar_length = int(weight * 30)
    bar = "█" * bar_length + "░" * (30 - bar_length)
    print(f"  {tag:8s} [{bar}] {weight:.2f}")

print()

# ============================================================================
# 第六部分：向量相似度分析
# ============================================================================

print("🔍 第六部分：向量相似度分析")
print("-" * 80)

# 模拟向量相似度计算
import random
random.seed(42)

test_queries = [
    "夜晚城市追逐场面",
    "森林中的逃亡",
    "人物情绪特写"
]

print("\n搜索测试案例:")
print("-" * 40)

for query in test_queries:
    print(f"\n查询: \"{query}\"")
    print("  匹配结果:")
    
    # 模拟匹配结果
    for i, beat in enumerate(beats_data['beats']):
        # 简单的关键词匹配模拟
        score = 0.0
        query_words = query.split()
        beat_tags = beat['tags']
        
        for word in query_words:
            for tag in beat_tags:
                if word in tag or tag in word:
                    score += 0.3
        
        # 添加一些随机性
        score = min(1.0, score + random.uniform(-0.1, 0.1))
        
        if score > 0.3:
            bar_length = int(score * 20)
            bar = "█" * bar_length + "░" * (20 - bar_length)
            print(f"    {i+1}. {beat['title']:20s} [{bar}] {score:.2f}")

print()

# ============================================================================
# 第七部分：导出历史记录
# ============================================================================

print("📋 第七部分：导出历史记录")
print("-" * 80)

export_records = [
    {
        "type": "script_docx",
        "file": f"{project_data['id']}_script.docx",
        "size": docx_path.stat().st_size if 'docx_path' in locals() and docx_path.exists() else 0,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    },
    {
        "type": "script_pdf",
        "file": f"{project_data['id']}_script.pdf",
        "size": pdf_path.stat().st_size if 'pdf_path' in locals() and pdf_path.exists() else 0,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    },
    {
        "type": "beatboard_image",
        "file": f"{project_data['id']}_beatboard.png",
        "size": image_path.stat().st_size if 'image_path' in locals() and image_path.exists() else 0,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
]

print("\n导出记录:")
print("-" * 40)
print(f"{'类型':<20} {'文件名':<35} {'大小':<15} {'时间':<20}")
print("-" * 90)

for record in export_records:
    if record['size'] > 0:
        size_str = f"{record['size']:,} B ({record['size']/1024:.1f} KB)"
        print(f"{record['type']:<20} {record['file']:<35} {size_str:<15} {record['time']:<20}")

print()

# ============================================================================
# 第八部分：系统统计
# ============================================================================

print("📊 第八部分：系统统计")
print("-" * 80)

total_tags = sum(len(tags) for scenes in tag_hierarchy.values() for tags in scenes.values())
total_categories = len(tag_hierarchy)
total_beats = len(beats_data['beats'])
total_exports = sum(1 for r in export_records if r['size'] > 0)

print(f"\n项目统计:")
print(f"  • 项目名称: {project_data['title']}")
print(f"  • Beat数量: {total_beats}")
print(f"  • 标签类别: {total_categories}")
print(f"  • 总标签数: {total_tags}")
print(f"  • 成功导出: {total_exports} 个文件")

if 'output_dir' in locals():
    total_size = sum(f.stat().st_size for f in output_dir.glob('*') if f.is_file())
    print(f"  • 导出总大小: {total_size:,} 字节 ({total_size/1024:.1f} KB)")

print()

# ============================================================================
# 总结
# ============================================================================

print("=" * 80)
print("✅ MVP演示完成！")
print("=" * 80)
print()
print("已实现的功能:")
print("  ✓ 数据库Schema扩展（标签层级、导出历史）")
print("  ✓ 剧本文档导出（DOCX和PDF格式）")
print("  ✓ BeatBoard图片导出（PNG格式）")
print("  ✓ 标签层级可视化展示")
print("  ✓ 标签权重管理示例")
print("  ✓ 向量相似度搜索模拟")
print("  ✓ 导出历史记录管理")
print()
if 'output_dir' in locals():
    print(f"导出文件位置: {output_dir.absolute()}")
    print()
print("下一步:")
print("  1. 查看导出的文件（DOCX、PDF、PNG）")
print("  2. 在Web界面中集成这些功能")
print("  3. 在启动器中添加快捷按钮")
print("  4. 实现完整的标签管理界面")
print("  5. 集成真实的向量搜索API")
print()
print("=" * 80)
