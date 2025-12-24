"""
文档导出服务
支持DOCX和PDF格式的剧本导出
"""

from typing import Optional, Dict, Any
from pathlib import Path
import uuid
from datetime import datetime
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template

try:
    from weasyprint import HTML
except Exception:
    HTML = None

from database import Project, Beat, ExportHistory


class DocumentExporter:
    """文档导出服务"""
    
    def __init__(self, db: Session):
        self.db = db
        self.export_dir = Path("exports")
        self.export_dir.mkdir(exist_ok=True)
    
    async def export_script_docx(
        self,
        project_id: str,
        include_beats: bool = True,
        include_tags: bool = True,
        include_metadata: bool = True
    ) -> Dict[str, Any]:
        """导出剧本为DOCX格式"""
        try:
            # 加载项目数据
            project = self.db.query(Project).filter(Project.id == project_id).first()
            if not project:
                return {"status": "error", "message": "项目不存在"}
            
            # 加载Beat数据
            beats = []
            if include_beats:
                beats = self.db.query(Beat).filter(
                    Beat.project_id == project_id
                ).order_by(Beat.order_index).all()
            
            # 创建DOCX文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(project.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            if include_metadata:
                doc.add_paragraph(f"项目ID: {project.id}")
                doc.add_paragraph(f"创建时间: {project.created_at}")
                if project.logline:
                    doc.add_paragraph(f"一句话故事: {project.logline}")
                doc.add_paragraph(f"Beat数量: {len(beats)}")
                doc.add_paragraph()
            
            # 添加每个Beat
            for beat in beats:
                # Beat标题
                doc.add_heading(f"Beat {beat.order_index + 1}", level=1)
                
                # Beat内容
                if beat.content:
                    doc.add_paragraph(beat.content)
                    doc.add_paragraph()
                
                # 标签
                if include_tags:
                    doc.add_heading("标签", level=2)
                    
                    if beat.emotion_tags:
                        doc.add_paragraph(f"情绪: {', '.join(beat.emotion_tags)}")
                    if beat.scene_tags:
                        doc.add_paragraph(f"场景: {', '.join(beat.scene_tags)}")
                    if beat.action_tags:
                        doc.add_paragraph(f"动作: {', '.join(beat.action_tags)}")
                    if beat.cinematography_tags:
                        doc.add_paragraph(f"摄影: {', '.join(beat.cinematography_tags)}")
                    doc.add_paragraph()
                
                # 分页
                if beat != beats[-1]:
                    doc.add_page_break()
            
            # 保存文档
            filename = f"{project_id}_script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.export_dir / filename
            doc.save(str(file_path))
            
            # 记录导出历史
            export_record = ExportHistory(
                id=str(uuid.uuid4()),
                project_id=project_id,
                export_type="script_docx",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                file_format="docx",
                options={
                    "include_beats": include_beats,
                    "include_tags": include_tags,
                    "include_metadata": include_metadata
                },
                status="completed",
                created_at=datetime.now()
            )
            self.db.add(export_record)
            self.db.commit()
            
            return {
                "status": "success",
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "export_id": export_record.id
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"DOCX导出失败: {str(e)}"
            }
    
    async def export_script_pdf(
        self,
        project_id: str,
        template: str = "professional",
        include_beats: bool = True,
        include_tags: bool = True
    ) -> Dict[str, Any]:
        """导出剧本为PDF格式"""
        try:
            if HTML is None:
                return {"status": "error", "message": "PDF导出依赖未安装: weasyprint"}

            # 加载项目数据
            project = self.db.query(Project).filter(Project.id == project_id).first()
            if not project:
                return {"status": "error", "message": "项目不存在"}
            
            # 加载Beat数据
            beats = []
            if include_beats:
                beats = self.db.query(Beat).filter(
                    Beat.project_id == project_id
                ).order_by(Beat.order_index).all()
            
            # HTML模板
            html_template = """
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {
                        size: A4;
                        margin: 2cm;
                    }
                    body {
                        font-family: 'Microsoft YaHei', Arial, sans-serif;
                        line-height: 1.6;
                        color: #333;
                    }
                    h1 {
                        color: #fbbf24;
                        text-align: center;
                        font-size: 32px;
                        margin-bottom: 20px;
                    }
                    h2 {
                        color: #3b82f6;
                        font-size: 24px;
                        margin-top: 30px;
                        border-bottom: 2px solid #3b82f6;
                        padding-bottom: 5px;
                    }
                    h3 {
                        color: #10b981;
                        font-size: 18px;
                        margin-top: 20px;
                    }
                    .metadata {
                        background: #f3f4f6;
                        padding: 15px;
                        border-radius: 5px;
                        margin-bottom: 30px;
                    }
                    .beat {
                        page-break-after: always;
                        margin-bottom: 40px;
                    }
                    .beat:last-child {
                        page-break-after: auto;
                    }
                    .tags {
                        background: #fef3c7;
                        padding: 10px;
                        border-radius: 5px;
                        margin-top: 10px;
                    }
                </style>
            </head>
            <body>
                <h1>{{ title }}</h1>
                <div class="metadata">
                    <p><strong>项目ID:</strong> {{ project_id }}</p>
                    <p><strong>创建时间:</strong> {{ created_at }}</p>
                    {% if logline %}
                    <p><strong>一句话故事:</strong> {{ logline }}</p>
                    {% endif %}
                    <p><strong>Beat数量:</strong> {{ beat_count }}</p>
                </div>
                
                {% for beat in beats %}
                <div class="beat">
                    <h2>Beat {{ loop.index }}</h2>
                    <p>{{ beat.content }}</p>
                    
                    {% if include_tags %}
                    <h3>标签</h3>
                    <div class="tags">
                        {% if beat.emotion_tags %}
                        <p><strong>情绪:</strong> {{ beat.emotion_tags | join(', ') }}</p>
                        {% endif %}
                        {% if beat.scene_tags %}
                        <p><strong>场景:</strong> {{ beat.scene_tags | join(', ') }}</p>
                        {% endif %}
                        {% if beat.action_tags %}
                        <p><strong>动作:</strong> {{ beat.action_tags | join(', ') }}</p>
                        {% endif %}
                        {% if beat.cinematography_tags %}
                        <p><strong>摄影:</strong> {{ beat.cinematography_tags | join(', ') }}</p>
                        {% endif %}
                    </div>
                    {% endif %}
                </div>
                {% endfor %}
            </body>
            </html>
            """
            
            # 渲染HTML
            template_obj = Template(html_template)
            html_content = template_obj.render(
                title=project.title,
                project_id=project.id,
                created_at=project.created_at,
                logline=project.logline,
                beat_count=len(beats),
                beats=beats,
                include_tags=include_tags
            )
            
            # 生成PDF
            filename = f"{project_id}_script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.export_dir / filename
            HTML(string=html_content).write_pdf(str(file_path))
            
            # 记录导出历史
            export_record = ExportHistory(
                id=str(uuid.uuid4()),
                project_id=project_id,
                export_type="script_pdf",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                file_format="pdf",
                options={
                    "template": template,
                    "include_beats": include_beats,
                    "include_tags": include_tags
                },
                status="completed",
                created_at=datetime.now()
            )
            self.db.add(export_record)
            self.db.commit()
            
            return {
                "status": "success",
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "export_id": export_record.id
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"PDF导出失败: {str(e)}"
            }

    async def export_script_markdown(
        self,
        project_id: str,
        include_beats: bool = True,
        include_tags: bool = True
    ) -> Dict[str, Any]:
        """导出剧本为Markdown格式"""
        try:
            project = self.db.query(Project).filter(Project.id == project_id).first()
            if not project:
                return {"status": "error", "message": "Project not found"}
                
            beats = self.db.query(Beat).filter(
                Beat.project_id == project_id
            ).order_by(Beat.order_index).all()
            
            md_lines = []
            
            # Metadata
            md_lines.append(f"# {project.title}")
            md_lines.append(f"> **Created**: {project.created_at}")
            if project.logline:
                md_lines.append(f"> **Logline**: {project.logline}")
            md_lines.append("\n---")
            
            # Beats
            for i, beat in enumerate(beats):
                md_lines.append(f"\n## Beat {i+1}")
                if beat.content:
                    md_lines.append(beat.content)
                
                if include_tags and (beat.emotion_tags or beat.scene_tags):
                    tags = []
                    if beat.emotion_tags: tags.extend([f"#{t}" for t in beat.emotion_tags])
                    if beat.scene_tags: tags.extend([f"#{t}" for t in beat.scene_tags])
                    md_lines.append(f"\n`{' '.join(tags)}`")
                    
                md_lines.append("\n")
            
            # Save
            filename = f"{project_id}_script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            file_path = self.export_dir / filename
            
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(md_lines))
                
            # History
            export_record = ExportHistory(
                id=str(uuid.uuid4()),
                project_id=project_id,
                export_type="script_md",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                file_format="md",
                status="completed",
                created_at=datetime.now()
            )
            self.db.add(export_record)
            self.db.commit()
            
            return {
                "status": "success",
                "file_path": str(file_path),
                "export_id": export_record.id
            }
        except Exception as e:
            return {"status": "error", "message": f"Markdown Export Failed: {str(e)}"}

